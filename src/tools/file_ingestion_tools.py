"""
File Ingestion Tools - Extract content from various file formats
Tools for processing PDFs, DOCs, TXT files and making content available to agents
"""

import os
import tempfile
from typing import Dict, Any, List, Optional
import json
from datetime import datetime

# File processing libraries
try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

def extract_text_from_pdf(file_path: str) -> Dict[str, Any]:
    """
    Extract text content from PDF files
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Dictionary with extracted text and metadata
    """
    try:
        extracted_text = ""
        metadata = {}
        
        # Try pdfplumber first (better for complex layouts)
        if pdfplumber:
            with pdfplumber.open(file_path) as pdf:
                metadata = {
                    "pages": len(pdf.pages),
                    "title": pdf.metadata.get('Title', ''),
                    "author": pdf.metadata.get('Author', ''),
                    "creator": pdf.metadata.get('Creator', '')
                }
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += f"\n--- Page {page_num} ---\n{page_text}\n"
        
        # Fallback to PyPDF2
        elif PyPDF2:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = {
                    "pages": len(pdf_reader.pages),
                    "title": pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    "author": pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else ''
                }
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += f"\n--- Page {page_num} ---\n{page_text}\n"
        else:
            return {
                "status": "error",
                "error": "PDF processing libraries not available. Install PyPDF2 or pdfplumber."
            }
        
        return {
            "status": "success",
            "content": extracted_text.strip(),
            "metadata": metadata,
            "content_type": "pdf",
            "word_count": len(extracted_text.split()),
            "extraction_method": "pdfplumber" if pdfplumber else "PyPDF2"
        }
        
    except Exception as e:
        return {
            "status": "error",
            "error": f"Failed to extract PDF content: {str(e)}"
        }

def extract_text_from_docx(file_path: str) -> Dict[str, Any]:
    """
    Extract text content from DOCX files
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Dictionary with extracted text and metadata
    """
    try:
        if not DocxDocument:
            return {
                "status": "error",
                "error": "python-docx library not available. Install python-docx."
            }
        
        doc = DocxDocument(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())
        
        extracted_text = "\n\n".join(paragraphs)
        
        # Extract metadata
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "paragraphs": len(paragraphs),
            "tables": len(doc.tables)
        }
        
        # Extract tables if present
        table_content = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            table_content.append(table_data)
        
        return {
            "status": "success",
            "content": extracted_text,
            "metadata": metadata,
            "tables": table_content,
            "content_type": "docx",
            "word_count": len(extracted_text.split())
        }
        
    except Exception as e:
        return {
            "status": "error",
            "error": f"Failed to extract DOCX content: {str(e)}"
        }

def extract_text_from_txt(file_path: str) -> Dict[str, Any]:
    """
    Extract text content from TXT files
    
    Args:
        file_path: Path to the TXT file
        
    Returns:
        Dictionary with extracted text and metadata
    """
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        content = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                    used_encoding = encoding
                    break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            return {
                "status": "error",
                "error": "Could not decode text file with any common encoding"
            }
        
        # Get file stats
        file_stats = os.stat(file_path)
        metadata = {
            "encoding": used_encoding,
            "size_bytes": file_stats.st_size,
            "modified": datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
            "lines": len(content.split('\n'))
        }
        
        return {
            "status": "success",
            "content": content,
            "metadata": metadata,
            "content_type": "txt",
            "word_count": len(content.split())
        }
        
    except Exception as e:
        return {
            "status": "error",
            "error": f"Failed to extract TXT content: {str(e)}"
        }

def process_uploaded_file(file_path: str, filename: str) -> Dict[str, Any]:
    """
    Process an uploaded file and extract its content based on file type
    
    Args:
        file_path: Path to the uploaded file
        filename: Original filename
        
    Returns:
        Dictionary with extracted content and processing results
    """
    try:
        # Determine file type from extension
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext == '.pdf':
            result = extract_text_from_pdf(file_path)
        elif file_ext in ['.docx']:
            result = extract_text_from_docx(file_path)
        elif file_ext in ['.doc']:
            # For .doc files, suggest conversion to .docx
            return {
                "status": "error",
                "error": "Legacy .doc files not supported. Please convert to .docx format."
            }
        elif file_ext in ['.txt']:
            result = extract_text_from_txt(file_path)
        else:
            return {
                "status": "error",
                "error": f"Unsupported file type: {file_ext}. Supported types: .pdf, .docx, .txt"
            }
        
        if result["status"] == "success":
            result.update({
                "filename": filename,
                "file_extension": file_ext,
                "processed_at": datetime.now().isoformat()
            })
        
        return result
        
    except Exception as e:
        return {
            "status": "error",
            "error": f"Failed to process file {filename}: {str(e)}"
        }

def chunk_content_for_analysis(content: str) -> List[Dict[str, Any]]:
    """
    Split large content into manageable chunks for agent processing with default settings
    
    Args:
        content: Text content to chunk
        
    Returns:
        List of content chunks with metadata
    """
    try:
        chunk_size = 2000
        overlap = 200
        
        if len(content) <= chunk_size:
            return [{
                "chunk_id": 1,
                "content": content,
                "start_pos": 0,
                "end_pos": len(content),
                "word_count": len(content.split())
            }]
        
        chunks = []
        start = 0
        chunk_id = 1
        
        while start < len(content):
            end = min(start + chunk_size, len(content))
            
            # Try to find a good breaking point (sentence or paragraph)
            if end < len(content):
                # Look for sentence endings near the end
                for i in range(end - 100, end):
                    if i > start and content[i] in '.!?\n':
                        end = i + 1
                        break
            
            chunk_content = content[start:end].strip()
            if chunk_content:
                chunks.append({
                    "chunk_id": chunk_id,
                    "content": chunk_content,
                    "start_pos": start,
                    "end_pos": end,
                    "word_count": len(chunk_content.split())
                })
                chunk_id += 1
            
            start = max(start + 1, end - overlap)
        
        return chunks
        
    except Exception as e:
        return [{
            "chunk_id": 1,
            "content": content,
            "error": f"Chunking failed: {str(e)}",
            "word_count": len(content.split())
        }]

def analyze_content_structure(content: str) -> Dict[str, Any]:
    """
    Analyze the structure and characteristics of extracted content
    
    Args:
        content: Extracted text content
        
    Returns:
        Analysis of content structure
    """
    try:
        lines = content.split('\n')
        words = content.split()
        
        # Basic statistics
        stats = {
            "total_characters": len(content),
            "total_words": len(words),
            "total_lines": len(lines),
            "average_words_per_line": len(words) / len(lines) if lines else 0
        }
        
        # Identify potential sections/headings
        potential_headings = []
        for i, line in enumerate(lines):
            line = line.strip()
            if line:
                # Check for typical heading patterns
                if (len(line) < 100 and 
                    (line.isupper() or 
                     line.startswith(('Chapter', 'Section', 'Part', '1.', '2.', '3.')) or
                     line.endswith(':'))):
                    potential_headings.append({
                        "line_number": i + 1,
                        "text": line
                    })
        
        # Identify lists
        list_indicators = ['•', '-', '*', '1.', '2.', '3.', 'a)', 'b)', 'c)']
        list_items = []
        for i, line in enumerate(lines):
            line = line.strip()
            for indicator in list_indicators:
                if line.startswith(indicator):
                    list_items.append({
                        "line_number": i + 1,
                        "text": line
                    })
                    break
        
        return {
            "status": "success",
            "statistics": stats,
            "structure": {
                "potential_headings": potential_headings[:10],  # Limit to first 10
                "list_items": list_items[:20],  # Limit to first 20
                "has_tables": "---" in content or "|" in content,
                "has_code": "```" in content or "def " in content or "function " in content
            }
        }
        
    except Exception as e:
        return {
            "status": "error",
            "error": f"Content analysis failed: {str(e)}"
        }
